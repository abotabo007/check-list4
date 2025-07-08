from flask import Blueprint, send_file, abort, session, render_template
from flask import current_app as app
from flask_login import login_required
from io import BytesIO
from docx import Document
from fleets.db import get_db
from fleets.helpers import permissions_required
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo



from flask import (
    Blueprint, flash, g, redirect, render_template, request, session, url_for, jsonify
)
from werkzeug.security import check_password_hash, generate_password_hash
from fleets.helpers import feedback, login_required, check_password, as_dict, permissions_required, check_inputs, best_fit, to_dict
from fleets.dictionaries import c1

from fleets.db import get_db


MAX_INSPECTIONS = 8

bp = Blueprint('views', __name__,)

@bp.before_app_request
def load_logged_in_user():
    user_id = session.get('user_id')

    if user_id is None:
        g.user = None
    else:
        g.user = get_db().execute(
            'SELECT * FROM users WHERE u_id = ?', (user_id,)
        ).fetchone()

# Handle requests to the root
@bp.route("/")
def index():
    """Handle requests to index"""
    # If user si not logged in, show an index page with some text
    if not session.get("user_id"):
        return render_template("index.html")

    # Get the user information from the DB
    users = []
    db = get_db()
    user = as_dict(db.execute("SELECT * FROM users WHERE u_id = ? AND c_id = ?", [session.get("user_id"), session.get("c_id")]).fetchall())
    vehicles = []

    # If the user is admin or owner give them the user list and the vehicle list, otherwise just give them just the vehicles
    if user:
        if user[0]["role"] in ["owner", "admin"]:
            users = as_dict(db.execute("SELECT * FROM users WHERE c_id = ?", [user[0]["c_id"]]).fetchall())
        vehicles = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY (number + 0)", [session["c_id"]]).fetchall())
    else:
        return redirect(url_for('.logout'))


    # Render index.html with the values from DB
    return render_template("index.html", user=user, users=users, vehicles=vehicles)

@bp.route("/register", methods=["GET", "POST"])
def register():
    """Register a new user"""
    # If request is a GET just show register.html
    if request.method == "GET":
        return render_template("register.html")

    # If request is post, process the data
    else:
        # Check that all inputs have data, if not, render an apology
        checks = check_inputs(request.form)
        if checks[0]:
            return feedback(checks[1].capitalize() + " field cannot be empty", "register.html", to_dict(request.form), "r")

        # Set the form data to variables
        name = request.form.get("username")
        email = request.form.get("email")
        company = request.form.get("company")

        # Get all the companies from DB
        db = get_db()
        companys = as_dict(db.execute("SELECT * FROM companys WHERE name = ?", [company]).fetchall())
        #

        # If company name exists return apology
        if len(companys) > 0:
            return feedback("Company name already exists", "register.html", to_dict(request.form), "r")

        # If password doesn't meet requierements return apology
        if check_password(request.form.get("password")):
            return feedback("Password does not meet requirements", "register.html", to_dict(request.form), "r")

        # If password and confirmation don't match return apology
        if request.form.get("password") != request.form.get("confirmation"):
            return feedback("Password doesn't match confirmation", "register.html", to_dict(request.form), "r")

        # Insert company and user into database with current user as owner
        hashed_password = generate_password_hash(request.form.get("password"))
        try:
            with db:
                cid = db.execute("INSERT INTO companys (name, owner) VALUES(?, ?)", (company, name)).lastrowid
                db.execute("INSERT INTO users (c_id, username, email, hash, role) VALUES(?, ?, ?, ?, ?)",
                            (cid, name, email, hashed_password, "owner"))
        except db.IntegrityError:
            # If there was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error registering Company/User, contact support', "register.html", to_dict(request.form), "r")
        else:
            # Flask confirmation and redirect to login page
            flash('Company/User registered')
            return redirect(url_for(".index"))

@bp.route("/login", methods=["GET", "POST"])
def login():
    """Log user in"""

    # Forget any user_id
    session.clear()

    # User reached route via POST (as by submitting a form via POST)
    if request.method == "POST":

        # Check that all inputs have data, if not, render an apology
        checks = check_inputs(request.form)
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "login.html", to_dict(request.form), "l", code=403)

        # Get company from database
        db = get_db()
        company = as_dict(db.execute("SELECT * FROM companys WHERE name = ?", [request.form.get("company")]).fetchall())

        # If company is not in db render an apology
        if len(company) != 1:
            return feedback("Company not found", "login.html", to_dict(request.form), "l", code=403)

        # Get user from DB
        rows = as_dict(db.execute("SELECT * FROM users WHERE username = ? AND c_id = ?", [request.form.get("username"), company[0]["id"]]).fetchall())

        # Ensure username exists and password is correct
        if len(rows) != 1 or not check_password_hash(rows[0]["hash"], request.form.get("password")):
            return feedback("Invalid username and/or password", "login.html", to_dict(request.form), "l", code=403)

        # Remember which user has logged in along with company and role
        session["user_id"] = rows[0]["u_id"]
        session["c_id"] = rows[0]["c_id"]
        session["role"] = rows[0]["role"]

        # Redirect user to home page
        return redirect(url_for("views.index"))

    # User reached route via GET (as by clicking a link or via redirect)
    else:
        return render_template("login.html")

@bp.route("/logout")
def logout():
    """Log user out"""

    # Forget any user_id
    session.clear()

    # Redirect user home
    flash("You logged out")
    return redirect(url_for('.index'))

@bp.route("/add-vehicle", methods=["GET", "POST"])
@login_required
@permissions_required
def add_vehicle():
    """Adds a new vehicle to the company"""
    # If request is GET render the page
    if request.method == "GET":
        return render_template("add-vehicle.html")

    # If request is post:
    else:
        # Check that all inputs have data (except tag), if not, render an apology
        checks = check_inputs(request.form, ["tag"])
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "add-vehicle.html", to_dict(request.form), "v")

        # Check if year is a 4digit number, if not, render an apology
        try:
            if int(request.form.get("year")) < 1900:
                return feedback("Year must have 4 digits", "add-vehicle.html", to_dict(request.form), "v")
        except:
            return feedback("Year must be a 4 digits number", "add-vehicle.html", to_dict(request.form), "v")

        # Create an array with the form data
        vehicle = [ session.get("c_id"),
                    request.form.get("make"),
                    request.form.get("model"),
                    request.form.get("year"),
                    request.form.get("number"),
                    request.form.get("vin"),
                    request.form.get("tag")]

        # Ensure Vehicle id is unique
        db = get_db()
        v = db.execute("SELECT * FROM vehicles WHERE number = ? AND c_id = ?", [request.form.get("number"), session.get("c_id")]).fetchall()
        if len(v) > 0:
            return feedback("A vehicle with that ID already exists in the company database", "add-vehicle.html", to_dict(request.form), "v")

        # Ensure VIN is unique
        v = db.execute("SELECT * FROM vehicles WHERE vin = ?", [request.form.get("vin")]).fetchall()
        if len(v) > 0:
            return feedback("A vehicle with that VIN already exists in the company database", "add-vehicle.html", to_dict(request.form), "v")

        # Insert Vehicle into database
        try:
            with db:
                db.execute("INSERT INTO vehicles (c_id, make, model, year, number, vin, tag) VALUES(?, ?, ?, ?, ?, ?, ?)", vehicle)
        except db.IntegrityError:
            # If there was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error adding vehicle, contact support', "add-vehicle.html", to_dict(request.form), "v")
        else:
            # Redirect to home
            flash("Vehicle added!")
            return redirect(url_for(".index"))

@bp.route("/add-user", methods=["GET", "POST"])
@login_required
@permissions_required
def add_user():
    """Adds a new user to the company"""
    # If request is GET render the page
    if request.method == "GET":
        return render_template("add-user.html")

    # If request is post:
    else:
        # Check that all inputs have data, if not, render an apology
        checks = check_inputs(request.form)
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "add-user.html", to_dict(request.form), "u")

        # If password doesn't meet requierements return apology
        if check_password(request.form.get("password")):
            return feedback("Password does not meet requirements", "add-user.html", to_dict(request.form), "u")

        # If password and confirmation don't match return apology
        if request.form.get("password") != request.form.get("confirmation"):
            return feedback("Passwords don't match", "add-user.html", to_dict(request.form), "u")

        # If role is not admin or user render an apology (Can only have 1 owner)
        if request.form.get("role") not in ["admin", "user"]:
            return feedback("Wrong role", "add-user.html", to_dict(request.form), "u")

        # Set the data into an array
        user = [ session.get("c_id"),
                    request.form.get("username"),
                    request.form.get("email"),
                    generate_password_hash(request.form.get("password")),
                    request.form.get("role")]

        # Query DB for the same username or email in the company, if exists return an apology
        db = get_db()
        v = db.execute("SELECT * FROM users WHERE (username = ? OR email = ?) AND c_id = ?",
                         [request.form.get("username"), request.form.get("email"), session.get("c_id")]).fetchall()
        if len(v) > 0:
            return feedback("Username/Email already exists in the company database", "add-user.html", to_dict(request.form), "u")

        # Insert user into DB
        try:
            with db:
                db.execute("INSERT INTO users (c_id, username, email, hash, role) VALUES(?, ?, ?, ?, ?)", user)
        except db.IntegrityError:
            # If there was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error adding user, contact support', "add-user.html", to_dict(request.form), "u")
        else:
            # Redirect to home

            flash("User added!")
            return redirect(url_for(".index"))

@bp.route("/inspection", methods=["GET", "POST"])
# @login_required
def inspection():
    """Creates an inspection"""

    db = get_db()

    if request.method == "GET":
        # ✅ Modifica indispensabile: usa c_id di default se non loggato
        c_id = session.get("c_id", 1)   

        # If vehicle is not in get request
        if not request.args.get("vehicle"):
            # Get all the company vehicles from DB
            vehicles = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY number", [c_id]).fetchall())

            # If only one vehicle just go to inspect that one, otherwise render template to select it
            if len(vehicles) == 1:
                return redirect(url_for(".inspection") + "?vehicle=" + str(vehicles[0]["number"]))
            else:
                return render_template("inspection.html", vehicles=vehicles)

        # If vehicle is in get request
        else:
            # ✅ Modifica indispensabile: usa c_id anche qui
            v = as_dict(db.execute("SELECT * FROM vehicles WHERE number = ? AND c_id = ?",
                                    [request.args.get("vehicle"), c_id]).fetchall())

            # If no such vehicle redirect to inspection
            if len(v) != 1:
                flash('Wrong parameter', 'error')
                return redirect(url_for(".inspection"))

            # Get next oil change from last inspection
            oil = db.execute("SELECT next_oil FROM inspections WHERE v_id = ? ORDER BY date DESC", [v[0]["v_id"]]).fetchone()

            # pass the info to render the inspection for that vehicle
            return render_template("inspection.html", inspection=c1, vehicle=request.args.get("vehicle"),
                                        v=v[0], oil=oil, date = datetime.now(ZoneInfo("Europe/Rome")).strftime('%d/%m/%Y %H:%M:%S'))

    # If request is post (Meaning: user submitted an inspection)
    else:
        # Query DB for that vehicle
        v = as_dict(db.execute("SELECT * FROM vehicles WHERE v_id = ? AND c_id = ?",
                                [request.form.get("vehicle"), session.get("c_id")]).fetchall())
        if v:
            oil = db.execute("SELECT next_oil FROM inspections WHERE v_id = ? ORDER BY date DESC", [v[0]["v_id"]]).fetchone()
        else:
            flash('Something went wrong with that request')
            return redirect(url_for('.inspection'))

        # Check that the important inputs have data, if not, render an apology
        checks = check_inputs(request.form, ["vehicle", "miles", "maintenance", "first_name", "last_name"], False)
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "inspection.html", to_dict(request.form), "i", 400,
                            inspection=c1, vehicle=v[0]["number"], v=v[0], oil=request.form.get("maintenance"),
                            date = datetime.now(ZoneInfo("Europe/Rome")).strftime('%d/%m/%Y %H:%M:%S'))

        # Check inspection is complete
        for c in c1:
            check = check_inputs(request.form, [c[0]], False)
            if check[0]:
                return feedback("Must provide value for: " + c[3].capitalize(), "inspection.html", to_dict(request.form), "i", 400,
                                inspection=c1, vehicle=v[0]["number"], v=v[0], oil=request.form.get("maintenance"),
                                date = datetime.now(ZoneInfo("Europe/Rome")).strftime('%d/%m/%Y %H:%M:%S'))

        # Create the DB query dynamically depending on the data submitted
        query = "INSERT INTO inspections (c_id, u_id, v_id, miles, next_oil, date, first_name, last_name"
        values = "(?, ?, ?, ?, ?, ?, ?, ?"
        vars = [
            session.get("c_id"),
            session.get("user_id"),
            request.form.get("vehicle"),
            request.form.get("miles"),
            request.form.get("maintenance"),
            datetime.now(ZoneInfo("Europe/Rome")).strftime('%Y-%m-%d %H:%M:%S'),
            request.form.get("first_name"),
            request.form.get("last_name"),
        ]

        # iterate over c1 to get whatever the user submitted
        for c in c1:
            if request.form.get(c[0]):
                query += ", " + c[0]
                values += ", ?"
                vars.append(request.form.get(c[0]))
            if request.form.get(c[1]):
                query += ", " + c[1]
                values += ", ?"
                vars.append(request.form.get(c[1]))
            if request.form.get(c[2]):
                query += ", " + c[2]
                values += ", ?"
                vars.append(request.form.get(c[2]))

        # assemble the final query string
        query += ") VALUES " + values + ")"

        # Submit the query to the database
        try:
            with db:
                db.execute(query, vars)
        except db.IntegrityError:
            return feedback('Database: Error adding inspection, contact support', "inspection.html", to_dict(request.form), "i", 400,
                                inspection=c1, vehicle=v[0]["number"], v=v[0], oil=request.form.get("maintenance"),
                                date = datetime.now(ZoneInfo("Europe/Rome")).strftime('%d/%m/%Y %H:%M:%S'))
        else:
            flash("Inspection added!")
            return redirect(url_for(".index"))


@bp.route("/password", methods=["GET", "POST"])
@login_required
def password():
    """Change Password"""
    # If request is GET render the page
    if request.method == "GET":
        return render_template("password.html")

    # If request is post (Meaning: user sumbited a password change)
    else:
        # query DB for user
        db = get_db()
        user = as_dict(db.execute("SELECT * FROM users WHERE u_id = ?", [session.get("user_id")]).fetchall())


        # Check that all inputs have data, if not, render an apology
        checks = check_inputs(request.form)
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "password.html")

        # If password doesn't meet requierements return apology
        if check_password(request.form.get("password")):
            return feedback("Password does not meet requirements", "password.html")

        # Ensure username exists and password is correct
        if len(user) != 1 or not check_password_hash(user[0]["hash"], request.form.get("old-password")):
            return feedback("Wrong password", "password.html")

        # If password and confirmation don't match return apology
        if request.form.get("password") != request.form.get("confirmation"):
            return feedback("Password and confimation don't match", "password.html")

        # Update database with new password
        hashed_password = generate_password_hash(request.form.get("password"))
        db = get_db()
        try:
            with db:
                db.execute("UPDATE users SET hash = ? WHERE u_id = ?", [hashed_password, session.get("user_id")])
        except db.IntegrityError:
            # If there was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error changing password, contact support', "password.html")

        else:
            # Flask confirmation and redirect to home

            flash("Password changed!")
            return redirect(url_for(".index"))


@bp.route("/edit-vehicle", methods=["GET", "POST"])
@login_required
def edit_vehicle():
    """Edit a Vehicle already in the database"""
    # If request is GET render the page
    if request.method == "GET":
        # If vehicle is not in get request
        if not request.args.get("vehicle"):
            # Get all the company vehicles from DB
            db = get_db()
            vehicles = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY (number + 0)", [session.get("c_id")]).fetchall())

            # If only one vehicle just go to edit that one, otherwise render template to select it
            if len(vehicles) == 1:
                return redirect(url_for(".edit_vehicle") + "?vehicle=" + str(vehicles[0]["number"]))
            else:
                return render_template("edit-vehicle.html", vehicles=vehicles)

        # If vechicle is in get request
        else:
            # Query DB for that vehicle
            db = get_db()
            v = as_dict(db.execute("SELECT * FROM vehicles WHERE number = ? AND c_id = ?",
                                    [request.args.get("vehicle"), session.get("c_id")]).fetchall())


            # If no such vehicle redirect to edit-vehicle otherwise pass the info to render the edit template for that vehicle
            if len(v) == 0:
                return redirect('/edit-vehicle')
            else:
                return render_template("edit-vehicle.html", vehicle=request.args.get("vehicle"), v=v[0])

    # If request is post (Meaning: user sumbited an edit)
    else:
        # Check that all the inputs have data except for tag, if not, render an apology
        checks = check_inputs(request.form, ["tag"])
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "edit-vehicle.html", to_dict(request.form), "v",
                                400, vehicle=request.form.get("vehicle"))

        # Check if year is a 4digit number, if not, render an apology
        try:
            if int(request.form.get("year")) < 1900:
                return feedback("Year must have 4 digits", "edit-vehicle.html", to_dict(request.form), "v", 400,
                                    vehicle=request.form.get("vehicle"))
        except:
            return feedback("Year must be a 4 digits number", "edit-vehicle.html", to_dict(request.form), "v",
                                400, vehicle=request.form.get("vehicle"))

        # Create an array with the form data
        vehicle = [ request.form.get("make"),
                    request.form.get("model"),
                    request.form.get("year"),
                    request.form.get("number"),
                    request.form.get("tag") ]

        # Get the vehicle from database
        db = get_db()
        v_id = db.execute("SELECT v_id FROM vehicles WHERE number = ? AND c_id = ?",
                            [request.form.get("vehicle"), session.get("c_id")]).fetchone()

        # If no vehicle, return an apology
        if not v_id:
            return feedback("Something went wrong with that request", "edit-vehicle.html", to_dict(request.form),
                                "v", 400, vehicle=request.form.get("vehicle"))

        # Ensure Vehicle id is unique
        v = db.execute("SELECT * FROM vehicles WHERE number = ? AND c_id = ?", [request.form.get("number"), session.get("c_id")]).fetchall()
        if len(v) > 0:
            return feedback("A vehicle with that ID already exists in the company database", "edit-vehicle.html", to_dict(request.form), "v")

        # Ensure VIN is unique
        v = db.execute("SELECT * FROM vehicles WHERE vin = ?", [request.form.get("vin")]).fetchall()
        if len(v) > 0:
            return feedback("A vehicle with that VIN already exists in the company database", "add-vehicle.html", to_dict(request.form), "v")

        # Add vehicle id to the vehicle array and update the vehicle with the new data
        vehicle.append(v_id[0])
        try:
            with db:
                db.execute("UPDATE vehicles SET make = ?, model = ?, year = ?, number = ?, tag = ? WHERE v_id = ?", vehicle)
        except db.IntegrityError:
            # If the was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error editing vehicle, contact support', "edit-vehicle.html",
                                to_dict(request.form), "v", 400, vehicle=request.form.get("vehicle"))
        else:
            # Flash the user the confimation and redirect to home

            flash('Database Updated!')
            return redirect(url_for(".index"))


@bp.route("/edit-user", methods=["GET", "POST"])
@permissions_required
@login_required
def edit_user():
    """Edit a user already in the database"""
    # If request is GET render the page
    if request.method == "GET":
        # If user is not in get request
        if not request.args.get("user"):
            # Get all the company users from DB except for the owner
            db = get_db()
            users = as_dict(db.execute('''SELECT * FROM users WHERE c_id = ? AND role != "owner"''', [session.get("c_id")]).fetchall())

            # If only one user just go to edit that one, otherwise render template to select it
            if len(users) == 1:
                return redirect(url_for(".edit_user") + "?user=" + str(users[0]["username"]))
            else:
                return render_template("edit-user.html", users=users)

        # If user is in get request
        else:
            # Query DB for that user and the company
            db = get_db()
            u = as_dict(db.execute("SELECT * FROM users WHERE username = ? AND c_id = ?", [request.args.get("user"), session.get("c_id")]).fetchall())
            c = as_dict(db.execute("SELECT * FROM companys WHERE id = ?", [session.get("c_id")]).fetchall())
            users = as_dict(db.execute('''SELECT * FROM users WHERE c_id = ? AND role != "owner"''', [session.get("c_id")]).fetchall())


            # If trying to edit owner return an apology
            if c[0]["owner"] == request.args.get("user"):
                return feedback("Can't edit the owner", "edit-user.html", users, "users")

            # If no such user redirect to edit-user otherwise pass the info to render the edit template for that user
            if len(u) != 1:
                return feedback('No such user exists in the company', "edit-user.html", users, "users")
            else:
                return render_template("edit-user.html", user=request.args.get("user"), u=u[0])

    # If request is post (Meaning: user sumbited an edit)
    else:
        # Get the user and all other users from the DB
        db = get_db()
        u = as_dict(db.execute("SELECT * FROM users WHERE u_id = ?", [request.form.get("u")]).fetchall())
        others = as_dict(db.execute("SELECT * FROM users WHERE u_id != ? AND c_id = ?",
                            [request.form.get("u"), session.get("c_id")]).fetchall())


        # Check that all the inputs have data, if not, render an apology
        checks = check_inputs(request.form)
        if checks[0]:
            return feedback("Must provide " + checks[1].capitalize(), "edit-user.html", u[0], 'u',
                                400, user=u[0]["username"])

        # If user is not in DB return an apology
        if len(u) != 1:
            return feedback("Something went wrong with that request", "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))

        # If trying to submit a user/email that already exists return an apology
        for other in others:
            if request.form.get("username") == other["username"] or request.form.get("email") == other["email"]:
                return feedback("Username/email already in company database", "edit-user.html", to_dict(request.form), 'u',
                                    400, user=request.form.get("u"))

        # If trying to edit the role of the owner return an apology
        if u[0]["role"] == "owner" and request.form.get("role") != "owner":
            return feedback("Can't change role of the owner", "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))

        # If trying to sumbit a role not supported return an apology
        if request.form.get("role") not in ["admin", "user"]:
            return feedback("Wrong role", "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))

        # If password doesn't meet requierements return apology
        if check_password(request.form.get("password")):
            return feedback("Password does not meet requirements", "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))

        # If password and confirmation don't match return apology
        if request.form.get("password") != request.form.get("confirmation"):
            return feedback("Password doesn't match confirmation", "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))

        # Set the data into an array
        user = [ request.form.get("username"),
                    request.form.get("email"),
                    generate_password_hash(request.form.get("password")),
                    request.form.get("role"),
                    request.form.get("u")]

        # Insert user into DB
        db = get_db()
        try:
            with db:
                db.execute("UPDATE users SET username = ?, email = ?, hash = ?, role = ? WHERE u_id = ?", user)
        except db.IntegrityError:
            # If there was an error flash the user
            print(db.IntegrityError)
            return feedback('Database: Error editing user, contact support', "edit-user.html", to_dict(request.form), 'u',
                                400, user=request.form.get("u"))
        else:
            # Redirect to home

            flash('Database Updated!')
            return redirect(url_for(".index"))


@bp.route("/vehicles", methods=["GET", "POST"])
@permissions_required
@login_required
def vehicles():
    """View Vehicles"""

    def get_inspections(request_arg):
        """Get inspections from database"""
        # Get the vehicle from DB
        db = get_db()
        vehicle = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? AND number = ?",
                                [session.get("c_id"), request_arg.get("vehicle")]).fetchall())

        # If vehicle submited not in database redirect to vehicles
        if len(vehicle) != 1:

            flash("Select a vehicle from the list", 'error')
            return redirect(url_for("views.vehicles"))

        # Get all the vehicles from DB
        v = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY (number + 0)",
                                [session.get("c_id")]).fetchall())

        # Get MAX_INSPECTIONS fron the vehicle sumbited and the users from DB
        inspections = as_dict(db.execute('''SELECT * FROM inspections, users WHERE inspections.c_id = ? AND inspections.v_id = ?
                                            AND inspections.u_id = users.u_id ORDER BY inspections.date DESC LIMIT ?''',
                                            [session.get("c_id"), vehicle[0]["v_id"], MAX_INSPECTIONS]).fetchall())


        # Create an array called inspection with this structure:
        # [Issue description, Issue name, Date, User (that made the inspection)] for every issue and inspection
        inspection = []

        # itearete over every inspection
        for i in inspections:
            found_issue = False

            #iterate over every issue description
            for c in c1:
                # If the inspection has this issue flagged append a new array
                if i[c[0]] == 0:
                    found_issue = True
                    inspection.append([i[c[1]], c[3], i["date"], i["username"]])

            # If there was no issue in the inspection we still want to show inspection data with no issue
            if not found_issue:
                inspection.append(["No issue", "No issue", i["date"], i["username"]])

        # If no inspections just create an array with no data to show in page
        if len(inspection) < 1:
            inspection = [["No data", "No data", "No data", "No data"]]

        # If request is get just pass the data to render the template, otherwise jsonify data for javascript
        if request.method == "GET":
            return render_template("vehicles.html", vehicle=request.args.get("vehicle"), inspection=inspection, vehicles=v)
        else:
            return jsonify(inspection)

    # If request is GET render the page
    if request.method == "GET":
        # If vehicle is not in get request
        if not request.args.get("vehicle"):
            # Get all the company vehicles and inspections from DB
            db = get_db()
            vehicles = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY (number + 0)", [session.get("c_id")]).fetchall())
            inspections = as_dict(db.execute("SELECT * FROM inspections WHERE c_id = ? ORDER BY date DESC", [session.get("c_id")]).fetchall())


            # Fancy way of creating a dictionary with the vehicles as keys and a list of inspections as values
            v = {ve["number"]:[[i["next_oil"], i["miles"], i["date"]] for i in inspections if i["v_id"] == ve["v_id"]] for ve in vehicles}

            # Append a projection of the date of the next oil change
            for vehicle, i in v.items():
                # If there are no inspections just append no data array
                if len(i) < 1:
                    i.append(["No data", "No data", "No data", "No data"])
                # If there is more than 2, we can proceed with the calculation
                elif len(i) > 2:
                    # Variable that we need for the best_fit function
                    miles_oil = 0
                    miles = []
                    dates = []
                    j = 0
                    # Need to convert dates to numbers to calculate best fit. Milliseconds from 01/01/1970 seem appropiate
                    d2 = datetime.strptime("1970-01-01", '%Y-%m-%d')
                    for array in i:
                        if j > MAX_INSPECTIONS:
                            break
                        # Get latest mileage to oil change
                        miles_oil = max(miles_oil, array[0])
                        # Append the mileage to the miles array
                        miles.append(array[1])
                        # Append the date in milliseconds from 1970 to the dates array
                        d1 = datetime.strptime(array[2], '%Y-%m-%d %H:%M:%S')
                        dates.append((d1 - d2)/timedelta(milliseconds=1))
                        j += 1
                    # Calculate projection
                    next_oil = d2 + timedelta(milliseconds=best_fit(dates, miles, miles_oil))
                    # Append calculation in date format
                    i[0].append(next_oil.strftime('%Y-%m-%d'))
                else:
                    i[0].append("Need more data")

            # Render template passing the data
            return render_template("vehicles.html", vehicles=v)

        # If vehicle is in get request, let the get_inspections function handle it
        else:
            return get_inspections(request.args)

    # If the request is post with the value of vehicle let the function handle the javascript request
    elif request.form:
        return get_inspections(request.form)

    # If the request is post without the value of vehicle
    else:

        # Get all the vehicles and inspections of the company
        db = get_db()
        inspections = as_dict(db.execute("SELECT v_id, date, miles, next_oil FROM inspections WHERE c_id = ? ORDER BY date DESC", [session.get("c_id")]).fetchall())
        vehicles = as_dict(db.execute("SELECT * FROM vehicles WHERE c_id = ? ORDER BY number", [session.get("c_id")]).fetchall())


        # Create the data for the graphs for all vehicles including projections
        graph_data = []
        # Iterate over every vehicle
        for v in vehicles:
            # Dictionary of dates and miles for the vehicle
            d = {v["number"]:{"dates":[],"miles":[]}}
            miles_oil = 0
            j = 0
            # Iterate over every inspection
            for i in inspections:
                # If current inspection corresponds to the current vehicle append date and miles of this inspection
                if i["v_id"] == v["v_id"]:
                    if j > MAX_INSPECTIONS:
                        break
                    d1 = datetime.strptime(i["date"], '%Y-%m-%d')
                    d2 = datetime.strptime("1970-01-01", '%Y-%m-%d')
                    d[v["number"]]["dates"].append((d1 - d2)/timedelta(milliseconds=1))
                    d[v["number"]]["miles"].append(i["miles"])
                    miles_oil = max(i["next_oil"], miles_oil)
                    j += 1
            d[v["number"]]["dates"].reverse()
            d[v["number"]]["miles"].reverse()
            # Calculate projections with al the inspections
            next_oil = best_fit(d[v["number"]]["dates"], d[v["number"]]["miles"], miles_oil)
            if next_oil != 0:
                d[v["number"]]["dates"].append(next_oil)
                d[v["number"]]["miles"].append(miles_oil)
            graph_data.append(d)

        return jsonify(graph_data)

@bp.route("/all_inspections")
@login_required
@permissions_required  # Only admin/owner can view all checklists
def all_inspections():
    db = get_db()
    inspections = as_dict(db.execute(
        """
        SELECT inspections.*, users.username, users.role, vehicles.vin, vehicles.number
        FROM inspections
        LEFT JOIN users ON inspections.u_id = users.u_id
        LEFT JOIN vehicles ON inspections.v_id = vehicles.v_id
        WHERE inspections.c_id = ?
        ORDER BY inspections.date DESC
        """,
        [session.get("c_id")]
    ).fetchall())
    # inside your route
    for ins in inspections:
        checklist_items = []
        for key, value in ins.items():
            if (
                key not in ['i_id', 'c_id', 'u_id', 'v_id', 'miles', 'next_oil', 'next_1', 'next_2', 'date', 'first_name', 'last_name', 'username', 'role']
                and not key.endswith('_t') and not key.endswith('_b')
            ):
                checklist_items.append((key, value))
        ins['checklist_items'] = checklist_items

    return render_template(
        "all_inspections.html",
        inspections=inspections
    )
@bp.route("/inspection/export/<int:inspection_id>")
@login_required
@permissions_required  # Only admin/owner can export checklists
def export_inspection(inspection_id):
    db = get_db()

    # Fetch inspection, joined user info, and VIN from vehicles
    inspection = db.execute("""
        SELECT inspections.*, users.username, vehicles.vin, vehicles.number, vehicles.tag
        FROM inspections
        LEFT JOIN users ON inspections.u_id = users.u_id
        LEFT JOIN vehicles ON inspections.v_id = vehicles.v_id
        WHERE inspections.i_id = ?
    """, (inspection_id,)).fetchone()

    if inspection is None:
        abort(404, description="Ispezione non trovata")

    # (Optional) Only allow the user who performed the inspection or admins/owners
    user_id = session.get("user_id")
    user_role = session.get("role", "")
    if user_id != inspection["u_id"] and user_role not in ["admin", "owner"]:
        abort(403, description="Accesso negato")

    # Prepare checklist fields for export (excluding metadata and text/blob fields)
    exclude_keys = {'i_id', 'c_id', 'u_id', 'v_id', 'miles', 'next_oil', 'next_1', 'next_2',
                    'date', 'first_name', 'last_name', 'username', 'role', 'vin'}
    checklist = []
    for key, value in dict(inspection).items():
        if key in exclude_keys or key.endswith('_t') or key.endswith('_b'):
            continue
        checklist.append((key, value))

    # SLICE checklist if ambulance
    if inspection['vin'] != 'ambulanza' and inspection['vin'] != 'mezzo_1':
        checklist = checklist[:13]

    doc = Document()
    doc.add_heading("Checklist Ispezione", level=1)

    # Add metadata
    doc.add_paragraph(f"Operatore: {inspection['first_name']} {inspection['last_name']}")
    doc.add_paragraph(f"Username: {inspection['username']}")
    doc.add_paragraph(f"Data ispezione: {inspection['date']}")
    doc.add_paragraph(f"Targa veicolo: {inspection['number']}")
    doc.add_paragraph(f"Sigla veicolo: {inspection['tag']}")
    doc.add_paragraph(f"Chilometraggio: {inspection['miles']}")
    doc.add_paragraph("")

    doc.add_heading("Risposte", level=2)
    for key, value in checklist:
        risposta = "OK" if value == 1 else "Non OK" if value == 0 else "N/D"
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {risposta}", style="List Bullet")

    # Prepare in-memory file
    fakefile = BytesIO()
    doc.save(fakefile)
    fakefile.seek(0)

    # Filename with date and user for clarity
    filename = f"ispezione_{inspection['date']}_{inspection['username']}.docx"

    return send_file(
        fakefile,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@bp.route("/vehicle/<int:vehicle_id>/qrcode")
@login_required
def vehicle_qrcode(vehicle_id):
    """
    Genera un QR code per un veicolo specifico e lo restituisce come immagine PNG.
    L'utente deve appartenere alla stessa azienda del veicolo.
    """
    db = get_db()
    # Verifica che il veicolo esista e appartenga alla company dell'utente
    vehicle = db.execute(
        "SELECT * FROM vehicles WHERE v_id = ? AND c_id = ?",
        [vehicle_id, session.get("c_id")]
    ).fetchone()
    if not vehicle:
        abort(404, description="Veicolo non trovato o accesso negato")

    # Dati che vuoi codificare nel QR (qui un URL, personalizza se vuoi)
    qr_data = f"{request.host_url.rstrip('/')}/inspection?vehicle={vehicle['number']}"

    # Generazione QR
    import qrcode
    import io
    img = qrcode.make(qr_data)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)

    # Nome file leggibile
    filename = f"qrcode_vehicle_{vehicle['number'] if 'number' in vehicle.keys() else vehicle_id}.png"


    return send_file(
        buf,
        mimetype='image/png',
        as_attachment=True,
        download_name=filename
    )